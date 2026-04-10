from docx import Document
import os

doc_path = r'd:\downloads\itlanta\69c0040b4a492_hackathon_problem_statement_v2_copy_20260323_150834 (3).docx'
doc = Document(doc_path)

print("=" * 80)
print("DOCUMENT CONTENT EXTRACTION")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'None'
        print(f"[{style}] {para.text}")

print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)

for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {r_idx}: " + " | ".join(cells))

print("\n" + "=" * 80)
print("DOCUMENT PROPERTIES")
print("=" * 80)
props = doc.core_properties
print(f"Title: {props.title}")
print(f"Author: {props.author}")
print(f"Created: {props.created}")
print(f"Modified: {props.modified}")
print(f"Subject: {props.subject}")
print(f"Category: {props.category}")

# Check for images
print(f"\nNumber of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")
print(f"Number of sections: {len(doc.sections)}")

# Check for inline shapes / images
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
print(f"Number of images: {image_count}")
